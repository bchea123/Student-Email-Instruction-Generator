# Importing the required libraries
import httplib2
import os
import gspread
import docx
from docx2pdf import convert
from oauth2client.service_account import ServiceAccountCredentials
from oauth2client.client import flow_from_clientsecrets
from oauth2client.file import Storage
from oauth2client.tools import run_flow

# Get current working directory
fileLocation = os.path.abspath(__file__)
cwd = os.path.dirname(fileLocation)

# Define the scope
scope = ['https://spreadsheets.google.com/feeds', 'https://www.googleapis.com/auth/drive']

# Add credentials to the account
creds = ServiceAccountCredentials.from_json_keyfile_name(cwd + '\\Read from Google Sheets-caccc118211a.json', scope)

# Client secret
client_secret = cwd + "\\client_secret.json"

# Storage
cred_storage = Storage('credentials.storage')

# Start the OAuth flow to retrieve credentials
def authorize_credentials():
    # Fetch credentials from storage
    credentials = cred_storage.get()
    # If the credentials don't exist in the storage location then run the flow
    if credentials is None or credentials.invalid:
        flow = flow_from_clientsecrets(client_secret, scope = scope)
        http = httplib2.Http()
        credentials = run_flow(flow, cred_storage, http = http)
    return credentials

credentials = authorize_credentials()

# Authorize the clientsheet
client = gspread.authorize(credentials)

# Get the instance of the Spreadsheet
sheet = client.open('Student Email Information')

# Get the first sheet of the Spreadsheet
sheet_instance = sheet.get_worksheet(0)


def setFirstNames(firstNames, rowStart, rowEnd):
    for i in range(rowStart, rowEnd):
        firstName = sheet_instance.cell(col = 3, row = i).value
        firstNames.append(firstName)

def setLastNames(lastNames, rowStart, rowEnd):
    for i in range(rowStart, rowEnd):
        lastName = sheet_instance.cell(col = 2, row = i).value
        lastNames.append(lastName)

def setEmails(emails, rowStart, rowEnd):
    for i in range(rowStart, rowEnd):
        email = sheet_instance.cell(col = 4, row = i).value
        emails.append(email)

def setPasswords(passwords, rowStart, rowEnd):
    for i in range(rowStart, rowEnd):
        password = sheet_instance.cell(col = 5, row = i).value
        passwords.append(password)

def writeDoc(firstNames, lastNames, emails, passwords, current, fileNames):
    doc = docx.Document(cwd + "\\Student Email Instructions (2-12-2021) - Copy.docx")

    #Add student name to title
    addTitle = doc.paragraphs[0].add_run(firstNames[current].upper() + " " + lastNames[current].upper())
    addTitle.font.name = "Calibri (Body)"
    addTitle.font.size = docx.shared.Pt(16)

    #Add student email to document
    addEmail = doc.paragraphs[5].add_run(emails[current])
    addEmail.font.name = "Calibri (Body)"
    addEmail.font.size = docx.shared.Pt(14)

    #Add email password to document
    addPassword = doc.paragraphs[6].add_run(passwords[current])
    addPassword.font.name = "Calibri (Body)"
    addPassword.font.size = docx.shared.Pt(14)

    fileName = cwd + "\\Student Email Instructions\\" + firstNames[current][0].lower() + lastNames[current].lower() + ".docx"

    doc.save(fileName)
    fileNames.append(fileName)

def contLoop():
    decision = input("Would you like to run again? (Y/N): ").upper()
    cont = False
    while True:
        if decision == "N":
            break
        elif decision == "Y":
            cont = True
            break
        else:
            decision = input("Invalid input. Please enter an appropriate response. (Y/N): ").upper()
    return cont

def main():
    while True:
        rowStart = int(input("Please enter the first row number: "))
        rowEnd = int(input("Please enter the last row number: ")) + 1

        print()
        print("=" * 75)

        firstNames = []
        lastNames = []
        emails = []
        passwords = []
        fileNames = []
        
        setFirstNames(firstNames, rowStart, rowEnd)
        setLastNames(lastNames, rowStart, rowEnd)
        setEmails(emails, rowStart, rowEnd)
        setPasswords(passwords, rowStart, rowEnd)

        print("First Names:", firstNames)
        print("=" * 75)
        print("Last Names:", lastNames)
        print("=" * 75)
        print("Emails:", emails)
        print("=" * 75)
        print("Passwords:", passwords)
        print("=" * 75)
        
        #Create Word documents based on queried information
        current = 0
        while current <= len(firstNames) - 1:
            writeDoc(firstNames, lastNames, emails, passwords, current, fileNames)
            current += 1

        print("FileNames:", fileNames)
        print("=" * 75)
        print()

        #Convert created Word documents into pdf files
        current = 0
        while current <= len(fileNames) - 1:
            newFileName = fileNames[current][0:len(fileNames[current]) - 4] + "pdf"
            convert(fileNames[current], newFileName)
            os.remove(fileNames[current])
            print()
            current += 1

        # Continue using application
        if contLoop() == False:
            break

        print()

main()